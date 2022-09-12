from docx import Document
import argparse
from libnmap.parser import NmapParser

parser = argparse.ArgumentParser(description='Convert nmap xml to docx notes')
parser.add_argument('-x', '--xml', type=str, help='Argument for nmap xml output file')
parser.add_argument('-d', '--docx', type=str, help='Argument for docx output')
parser.add_argument('-n', '--name', type=str, default='unnamed', help='Argument for project name')
arguments = parser.parse_args()

document = Document()
nmap_report = NmapParser.parse_fromfile(arguments.xml)


if arguments.xml and arguments.docx:
    document.add_heading(f'Notes for {arguments.name} project.', 0)
    for host in nmap_report.hosts:
        port = None
        state = None
        banner = None
        ip_addr = host.address
        document.add_heading(ip_addr, level=1)

        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Port'
        hdr_cells[1].text = 'State'
        hdr_cells[2].text = 'Service'

        for service in host.services:
            port = f'{service.port}/{service.protocol}'
            state = service.state
            banner = service.banner
            row_cells = table.add_row().cells
            row_cells[0].text = port
            row_cells[1].text = state
            row_cells[2].text = banner

        document.add_paragraph('\nNotes:\n')
        document.add_page_break()
    document.save(arguments.docx)
else:
    print('[-] - Input Error. Use --help for more help.')
